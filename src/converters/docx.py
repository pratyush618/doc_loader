from typing import Optional, BinaryIO
from docx import Document

from ..core.exceptions import ConversionException
from .base import BaseConverter, ConverterResult


class DocxConverter(BaseConverter):
    """Converter for Microsoft Word documents"""
    
    def accepts(self, file: BinaryIO, mimetype: Optional[str] = None,
                extension: Optional[str] = None) -> bool:
        """Check if this is a DOCX file"""
        if extension and extension.lower() in {'.docx', '.docm'}:
            return True
        
        if mimetype and mimetype in {
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.ms-word.document.macroEnabled.12'
        }:
            return True
        
        # Try to open as docx
        try:
            Document(file)
            return True
        except Exception:
            return False
        finally:
            self._reset_file_position(file)
    
    async def convert(self, file: BinaryIO, **kwargs) -> ConverterResult:
        """Convert DOCX to markdown"""
        try:
            doc = Document(file)
            
            # Convert content to markdown
            markdown_parts = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check for heading styles
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                        markdown_parts.append(f"{'#' * level} {para.text}\n")
                    else:
                        # Process runs for formatting
                        formatted_text = self._process_paragraph(para)
                        markdown_parts.append(f"{formatted_text}\n")
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                markdown_parts.append(self._convert_table(table))
            
            content = "\n".join(markdown_parts)
            
            # Extract title from first heading or paragraph
            title = None
            if content.strip():
                lines = content.split('\n')
                for line in lines:
                    clean_line = line.strip().replace('#', '').strip()
                    if clean_line:
                        title = clean_line[:100]  # First 100 chars
                        break
            
            return ConverterResult(
                content=content,
                title=title
            )
            
        except Exception as e:
            raise ConversionException(f"Failed to convert DOCX: {str(e)}")
    
    def _process_paragraph(self, paragraph) -> str:
        """Process paragraph with formatting"""
        result = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Apply formatting
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"_{text}_"
            
            result.append(text)
        
        return "".join(result)
    
    def _convert_table(self, table) -> str:
        """Convert table to markdown"""
        if not table.rows:
            return ""
        
        markdown_lines = []
        
        # Process header row
        header_cells = []
        for cell in table.rows[0].cells:
            header_cells.append(cell.text.strip())
        
        markdown_lines.append("| " + " | ".join(header_cells) + " |")
        markdown_lines.append("| " + " | ".join(["-" * len(cell) for cell in header_cells]) + " |")
        
        # Process data rows
        for row in table.rows[1:]:
            row_cells = []
            for cell in row.cells:
                row_cells.append(cell.text.strip())
            markdown_lines.append("| " + " | ".join(row_cells) + " |")
        
        return "\n".join(markdown_lines) + "\n"